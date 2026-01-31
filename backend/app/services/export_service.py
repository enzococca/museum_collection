"""
Export Service for Museum Collection Analytics
Exports reports to PDF, Excel, and DOCX formats.
"""
import io
import pandas as pd
from typing import Dict, Any, List
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import xlsxwriter


class ExportService:
    """Service for exporting analytics reports to various formats."""

    def __init__(self, report_data: Dict[str, Any]):
        """Initialize with report data."""
        self.report = report_data
        self.generated_at = datetime.utcnow()

    def export_to_excel(self) -> io.BytesIO:
        """Export comprehensive report to Excel format."""
        output = io.BytesIO()
        workbook = xlsxwriter.Workbook(output, {'in_memory': True})

        # Define formats
        header_format = workbook.add_format({
            'bold': True,
            'bg_color': '#4a5568',
            'font_color': 'white',
            'border': 1,
            'align': 'center',
            'valign': 'vcenter'
        })
        cell_format = workbook.add_format({
            'border': 1,
            'align': 'left',
            'valign': 'vcenter'
        })
        number_format = workbook.add_format({
            'border': 1,
            'align': 'right',
            'valign': 'vcenter',
            'num_format': '#,##0'
        })
        percent_format = workbook.add_format({
            'border': 1,
            'align': 'right',
            'valign': 'vcenter',
            'num_format': '0.0%'
        })
        title_format = workbook.add_format({
            'bold': True,
            'font_size': 14,
            'font_color': '#2d3748'
        })

        # Summary Sheet
        self._create_summary_sheet(workbook, header_format, cell_format, number_format, title_format)

        # Distribution Sheets
        for var_name, dist_data in self.report.get('distributions', {}).items():
            self._create_distribution_sheet(workbook, var_name, dist_data, header_format, cell_format, number_format, percent_format, title_format)

        # Cross-tabulation Sheets
        for corr_name, corr_data in self.report.get('correlations', {}).items():
            self._create_crosstab_sheet(workbook, corr_name, corr_data, header_format, cell_format, number_format, title_format)

        # Collection Comparison Sheet
        if 'collection_comparison' in self.report:
            self._create_collection_comparison_sheet(workbook, header_format, cell_format, number_format, title_format)

        workbook.close()
        output.seek(0)
        return output

    def _create_summary_sheet(self, workbook, header_format, cell_format, number_format, title_format):
        """Create summary statistics sheet."""
        worksheet = workbook.add_worksheet('Summary')
        worksheet.set_column('A:A', 30)
        worksheet.set_column('B:B', 20)

        row = 0
        worksheet.write(row, 0, 'Museum Collection Analytics Report', title_format)
        row += 1
        worksheet.write(row, 0, f'Generated: {self.generated_at.strftime("%Y-%m-%d %H:%M UTC")}', cell_format)
        row += 3

        worksheet.write(row, 0, 'Summary Statistics', title_format)
        row += 1

        summary = self.report.get('summary', {})
        metrics = [
            ('Total Artifacts', summary.get('total', 0)),
            ('Collections', summary.get('collections', 0)),
            ('Object Types', summary.get('object_types', 0)),
            ('Materials', summary.get('materials', 0)),
            ('Chronological Periods', summary.get('chronologies', 0)),
            ('On Display', summary.get('on_display', 0)),
        ]

        worksheet.write(row, 0, 'Metric', header_format)
        worksheet.write(row, 1, 'Value', header_format)
        row += 1

        for metric, value in metrics:
            worksheet.write(row, 0, metric, cell_format)
            worksheet.write(row, 1, value, number_format)
            row += 1

    def _create_distribution_sheet(self, workbook, var_name, dist_data, header_format, cell_format, number_format, percent_format, title_format):
        """Create distribution analysis sheet."""
        sheet_name = var_name.replace('_', ' ').title()[:31]  # Excel sheet name limit
        worksheet = workbook.add_worksheet(sheet_name)
        worksheet.set_column('A:A', 40)
        worksheet.set_column('B:C', 15)

        row = 0
        worksheet.write(row, 0, f'Distribution: {var_name}', title_format)
        row += 2

        # Summary info
        worksheet.write(row, 0, f'Total unique values: {dist_data.get("unique_values", 0)}')
        row += 1
        worksheet.write(row, 0, f'Concentration level: {dist_data.get("concentration_level", "N/A")}')
        row += 2

        # Distribution table
        worksheet.write(row, 0, 'Value', header_format)
        worksheet.write(row, 1, 'Count', header_format)
        worksheet.write(row, 2, 'Percentage', header_format)
        row += 1

        for item in dist_data.get('distribution', []):
            worksheet.write(row, 0, str(item.get('value', '')), cell_format)
            worksheet.write(row, 1, item.get('count', 0), number_format)
            worksheet.write(row, 2, item.get('percentage', 0) / 100, percent_format)
            row += 1

    def _create_crosstab_sheet(self, workbook, corr_name, corr_data, header_format, cell_format, number_format, title_format):
        """Create cross-tabulation sheet."""
        sheet_name = corr_name.replace('_vs_', ' x ')[:31]
        worksheet = workbook.add_worksheet(sheet_name)

        crosstab = corr_data.get('crosstab', {})
        chi_square = corr_data.get('chi_square', {})

        row = 0
        worksheet.write(row, 0, f'Cross-tabulation: {corr_name}', title_format)
        row += 2

        # Chi-square results
        worksheet.write(row, 0, 'Statistical Test Results:', title_format)
        row += 1
        worksheet.write(row, 0, f'Chi-square: {chi_square.get("chi_square", "N/A")}')
        row += 1
        worksheet.write(row, 0, f'P-value: {chi_square.get("p_value", "N/A")}')
        row += 1
        worksheet.write(row, 0, f'Significance: {chi_square.get("significance", "N/A")}')
        row += 1
        worksheet.write(row, 0, f'Effect strength: {chi_square.get("strength", "N/A")}')
        row += 2

        # Cross-tabulation table
        rows = crosstab.get('rows', [])
        cols = crosstab.get('columns', [])
        counts = crosstab.get('counts', [])

        if rows and cols and counts:
            # Header row
            worksheet.write(row, 0, '', header_format)
            for col_idx, col_name in enumerate(cols):
                worksheet.write(row, col_idx + 1, str(col_name), header_format)
            row += 1

            # Data rows
            for row_idx, row_name in enumerate(rows):
                worksheet.write(row, 0, str(row_name), header_format)
                for col_idx, count in enumerate(counts[row_idx]):
                    worksheet.write(row, col_idx + 1, count, number_format)
                row += 1

    def _create_collection_comparison_sheet(self, workbook, header_format, cell_format, number_format, title_format):
        """Create collection comparison sheet."""
        worksheet = workbook.add_worksheet('Collection Comparison')
        worksheet.set_column('A:A', 25)
        worksheet.set_column('B:E', 20)

        comparison = self.report.get('collection_comparison', {})
        collections = comparison.get('collections', {})

        row = 0
        worksheet.write(row, 0, 'Collection Comparison', title_format)
        row += 2

        # Collection summary table
        col_names = list(collections.keys())
        worksheet.write(row, 0, 'Metric', header_format)
        for col_idx, col_name in enumerate(col_names):
            worksheet.write(row, col_idx + 1, col_name.title(), header_format)
        row += 1

        metrics = ['total', 'on_display_pct']
        metric_labels = {'total': 'Total Artifacts', 'on_display_pct': 'On Display (%)'}

        for metric in metrics:
            worksheet.write(row, 0, metric_labels.get(metric, metric), cell_format)
            for col_idx, col_name in enumerate(col_names):
                value = collections[col_name].get(metric, 0)
                worksheet.write(row, col_idx + 1, value, number_format)
            row += 1

        # Commonalities and differences
        row += 2
        worksheet.write(row, 0, 'Commonalities:', title_format)
        row += 1
        for item in comparison.get('commonalities', []):
            worksheet.write(row, 0, item, cell_format)
            row += 1

        row += 1
        worksheet.write(row, 0, 'Differences:', title_format)
        row += 1
        for item in comparison.get('differences', []):
            worksheet.write(row, 0, item, cell_format)
            row += 1

    def export_to_docx(self) -> io.BytesIO:
        """Export comprehensive report to Word document format."""
        document = Document()

        # Title
        title = document.add_heading('Museum Collection Analytics Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle with date
        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f'Generated: {self.generated_at.strftime("%Y-%m-%d %H:%M UTC")}')
        run.italic = True

        document.add_paragraph()

        # Main narrative
        if self.report.get('main_narrative'):
            self._add_markdown_content(document, self.report['main_narrative'])

        # Summary Statistics
        document.add_heading('Summary Statistics', level=1)
        summary = self.report.get('summary', {})
        self._add_summary_table(document, summary)

        # Distribution analyses
        for var_name, dist_data in self.report.get('distributions', {}).items():
            document.add_heading(f'{var_name.replace("_", " ").title()} Distribution', level=2)

            # Add narrative if available
            if dist_data.get('narrative'):
                document.add_paragraph(dist_data['narrative'])

            # Add table
            self._add_distribution_table(document, dist_data)

        # Correlation analyses
        document.add_heading('Statistical Correlations', level=1)
        for corr_name, corr_data in self.report.get('correlations', {}).items():
            chi = corr_data.get('chi_square', {})
            if chi.get('interpretation'):
                p = document.add_paragraph()
                p.add_run(f'{corr_name}: ').bold = True
                p.add_run(chi['interpretation'])

        # Collection comparison
        if 'collection_comparison' in self.report:
            document.add_heading('Collection Comparison', level=1)
            comparison = self.report['collection_comparison']
            if comparison.get('narrative'):
                self._add_markdown_content(document, comparison['narrative'])

        # Save to bytes
        output = io.BytesIO()
        document.save(output)
        output.seek(0)
        return output

    def _add_markdown_content(self, document, markdown_text: str):
        """Add markdown-formatted content to document."""
        lines = markdown_text.split('\n')
        for line in lines:
            if line.startswith('## '):
                document.add_heading(line[3:], level=2)
            elif line.startswith('**') and line.endswith('**'):
                p = document.add_paragraph()
                p.add_run(line.strip('*')).bold = True
            elif line.startswith('- '):
                document.add_paragraph(line[2:], style='List Bullet')
            elif line.strip():
                document.add_paragraph(line)

    def _add_summary_table(self, document, summary: Dict):
        """Add summary statistics table."""
        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Metric'
        header_cells[1].text = 'Value'
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True

        # Data
        metrics = [
            ('Total Artifacts', summary.get('total', 0)),
            ('Collections', summary.get('collections', 0)),
            ('Object Types', summary.get('object_types', 0)),
            ('Materials', summary.get('materials', 0)),
            ('Chronological Periods', summary.get('chronologies', 0)),
            ('On Display', summary.get('on_display', 0)),
        ]

        for metric, value in metrics:
            row = table.add_row().cells
            row[0].text = metric
            row[1].text = str(value)

    def _add_distribution_table(self, document, dist_data: Dict):
        """Add distribution table to document."""
        distribution = dist_data.get('distribution', [])[:10]  # Top 10

        if not distribution:
            document.add_paragraph('No data available.')
            return

        table = document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Value'
        header_cells[1].text = 'Count'
        header_cells[2].text = 'Percentage'
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True

        # Data
        for item in distribution:
            row = table.add_row().cells
            row[0].text = str(item.get('value', ''))
            row[1].text = str(item.get('count', 0))
            row[2].text = f"{item.get('percentage', 0):.1f}%"


def generate_excel_report(report_data: Dict) -> io.BytesIO:
    """Generate Excel report from analytics data."""
    service = ExportService(report_data)
    return service.export_to_excel()


def generate_docx_report(report_data: Dict) -> io.BytesIO:
    """Generate Word document report from analytics data."""
    service = ExportService(report_data)
    return service.export_to_docx()
